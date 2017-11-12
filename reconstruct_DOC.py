
# coding: utf-8

# # 从Docx文件中读取表格
# 

# 导入必要的运行库: 
# 
# * collections 是用来产生有序字典
# * cleaner是自己编写的序列清理库
# 

# In[1]:


import pandas as pd
from pandas import DataFrame, Series
from collections import OrderedDict # 有序字典
from docx import Document
import re
import os
from save_sequence import save_seq_to_csv # 自建包, 保存数据至csv
from cleaner import chaintype_clean,name_clean,seq_3letters_clean, seq_1letters_clean, ID_clean


# ## 输入参数
# 
# 此处需要输入路径和文件名
# 
# *  doc将保存整个word文件
# *  作为子程序调用时, 本文件中所输入的参数不起作用. 请在compare_seq.ipynb文件中设定

# In[2]:


if __name__=="__main__":
    input_path='demo'
    fname='AAA.docx'
    output_path='demo'


# ## 此处指定表格类型
# 
# 需要手工指定表格类型, 注意编号从0开始, 
# 每一个表格都需要指定. 

# In[3]:


if __name__=="__main__":
    table_catalog_dict={
        0: {"head": 1, "seqtype":'chain', "chaintype":'HeavyChain'},
        1: {"head": 1, "seqtype":'chain', "chaintype":'LightChain'},
        2: {"head": 2, "seqtype":'CDR', "chaintype":'HC'}, 
        3: {"head": 2, "seqtype":'CDR', "chaintype":'LC'}, 
        4: {"head": 1, "seqtype":'chain', "chaintype":'HeavyChain'},
        5: {"head": 1, "seqtype":'chain', "chaintype":'LightChain'},
    }


# ## 将docx中的表格转换成pandas的Dataframe
# 
# * doc:  docx的文件内容将存储在doc变量中. 
# * idx: 表格的编号, 在docx文件中表格的编号, 注意python是从0开始编号的
# * table_catalog_dict:  入前所述, 为标记表格形式与内容的字典
# 

# In[4]:


def table2pandas(doc,idx,table_catalog_dict):
    # 提取编号为idx的表格
    table=doc.tables[idx]
    
    # 确定表格实际的列数colNumber, 和行数rowNumber
    # 注意表头中很可能有融合的单元格, 因此不能从首行中取得列数
    colNumber=len(table.rows[-1].cells) # 首行可能有表头, 所以从末尾行取
    rowNumber=len(table.column_cells(0))
    
    # 从table_catalog_dict中取得表头所占据的行数
    head=table_catalog_dict[idx]["head"]
    # 真正开始有内容的表格行就是从表头之后的行开始的. 因为首行=0, 所以实际起始行=0+head
    rowStart=0+head

    # 从表头中提取标题行内容
    if head == 0: 
        # 无标题行, 以数字代替
        colums_title=range(colNumber)
    else: 
        # 当存在标题行是, 需要找到离真实内容最接近的表头行
        colums_title=[each_cell.text for each_cell in table.row_cells(head-1)]
        
    # 使用有序字典作为转换成DataFrame的前身. 
    table_dict=OrderedDict();
    
    #遍历每一列, 将每一列的内容做成一个list
    for colIdx in range(colNumber):
        # colum_data是一个list包含每一列的所有单元格
        colum_data=[each_cell.text for each_cell in table.column_cells(colIdx)]
        # 只需要其中从起始列开始的剩余部分, 因为之前是表头内容
        colum=colum_data[rowStart:]
        # 压入字典中, 感觉稍微有点奇怪, 但反正做出来的是dataframe
        table_dict[colums_title[colIdx]]=colum
    return DataFrame(table_dict)


# # 分别从不同类型的表格中提取
# 
# 总的策略是先将不同表格的dataframe同一整理成ID-sequence-chaintype的形式, 完成之后再分别进行清洗. 
# 
# ## 从chain表格中提取
# 
# 表格都已经转换成了dataframe, chain表格一般是两列, (也有变态的3列), 首列是ID, 然后是sequence内容. 

# In[5]:


def get_chain_dataframe(pd_dataframe,tableID,table_catalog_dict):
    # 假定首列, 即第0列是ID列, 第1列是sequence列. 其他如果还有列就先不管了. 
    # 万一有变态的特殊表格, 再另行调整吧
    ID_column=pd_dataframe.columns[0]
    seq_column=pd_dataframe.columns[1] # 万一再碰到诡异的表格, 可能需要指定位置
    
    # 只提取ID列和sequence列
    chain_dataframe=pd_dataframe.loc[:,ID_column:seq_column].copy() # 只要ID和序列两列表格, 万一有更多列不要!!!
    
#     # 将ID列设定成索引列. 如果需要提取数量行的话, 应当使用.iloc[行数]定位
#     chain_dataframe.set_index(ID_column, inplace=True)
    # 统一序列的列名称为SEQUENCE
#     chain_dataframe.rename(columns={seq_column:"SEQUENCE"},inplace=True)
    chain_dataframe.columns=["NAME","SEQUENCE"]
    # 在docx文件中, 一个表格就是同一种chaintype类型. 
    # 从table_catalog_dict字典的记录中, 取得chaintype
    chaintype=table_catalog_dict[tableID]["chaintype"]
    chain_dataframe['CHAINTYPE']=chaintype

    return chain_dataframe


# In[6]:


if __name__=="__main__":
    # 读取文件
    filename=os.path.join(input_path,fname)
    doc = Document(filename)
    pd_dataframe=table2pandas(doc,0,table_catalog_dict)
    chain_dataframe=get_chain_dataframe(pd_dataframe,0,table_catalog_dict);
    print(chain_dataframe.iloc[0:2])## 从CDR表格中提取


# ## 从CDR表格中提取
# 
# 典型的CDR表格是有融合表头的, 然后下面的表格正文部分有4列: 
# * 首列第0列是ID
# * 第1至3列分别是CDR1, CDR2, CDR3
# * 融合的表头中有标记HC或LC, 但在生成dataframe时已经抛弃了. 所以完整的类型需用从table_catalog_dict字典中取得

# In[7]:


def get_CDR_dataframe(pd_dataframe,tableID,table_catalog_dict):
    # 假定首列, 即第0列是ID列, 第1列是sequence列. 其他如果还有列就先不管了. 
    # 万一有变态的特殊表格, 再另行调整吧     
    ID_column=pd_dataframe.columns[0]
    chaintype_base=table_catalog_dict[tableID]["chaintype"]
    
    # CDR_part_list这个列表暂存ID-sequence的对应项. 
    # 注意这里将CDR1,CDR2,CDR3的项目依次串联在一起. 但暂时仅仅在chaintype中简单记录, 之后再清洗chaintype
    CDR_part_list=list();
    # 遍历CDR123所在的3列. 
    for i in range(1,4):
        CDR_part=DataFrame()
        seq_column=pd_dataframe.columns[i]
        # 将ID列和当前列合并成一个dataframe
        CDR_part=pd_dataframe.loc[:,[ID_column,seq_column]].copy()
        # 简单记录chaintype
        CDR_part['CHAINTYPE']=chaintype_base+seq_column
        # 重命名sequence所在列
        CDR_part.rename(columns={seq_column:"SEQUENCE"},inplace=True)
        # 将做好的dataframe压入list中暂存
        CDR_part_list.append(CDR_part)
    
    # 将list中的三个dataframe连接
    CDR_dataframe=pd.concat(CDR_part_list)
    # 重设ID列为索引列
#     CDR_dataframe.set_index(ID_column, inplace=True)
    CDR_dataframe.columns=["NAME","SEQUENCE",'CHAINTYPE']
    return CDR_dataframe


# In[8]:


if __name__=="__main__":
    pd_dataframe=table2pandas(doc,3,table_catalog_dict)
    CDR_dataframe=get_CDR_dataframe(pd_dataframe,3,table_catalog_dict)
    print(CDR_dataframe.iloc[0:4])


# ## 将提取的数据合并生成Dataframe

# In[9]:


def make_seq_dataframe(doc,table_catalog_dict):
    seq_dataframe_list=list()
    for tableID in range(len(doc.tables)):
        pd_dataframe=table2pandas(doc,tableID,table_catalog_dict)
        if table_catalog_dict[tableID]["seqtype"]=='chain':
            seq_dataframe_list.append(get_chain_dataframe(pd_dataframe,tableID,table_catalog_dict))
        elif table_catalog_dict[tableID]["seqtype"]=='CDR':
            seq_dataframe_list.append(get_CDR_dataframe(pd_dataframe,tableID,table_catalog_dict))
    seq_dataframe=pd.concat(seq_dataframe_list)
    return seq_dataframe


# In[10]:


if __name__=="__main__":
    seq_dataframe=make_seq_dataframe(doc,table_catalog_dict)
    print(seq_dataframe)


# # 提取序列中的数字作为ID

# ## 将带有数字的序列复制到ID列

# In[11]:


def copy_seq_with_num_to_ID(df,col_from,col_to):
    df[col_to]=""
    df[col_to]=df[col_from]
    return df


# In[12]:


if __name__=="__main__":
    seq_dataframe=make_seq_dataframe(doc,table_catalog_dict)
#     print(seq_dataframe.iloc[0:2,:])
    seq_dataframe=copy_seq_with_num_to_ID(seq_dataframe,"SEQUENCE","ID")
    # print(seq_dataframe)


# ## 清洗ID

# In[13]:


if __name__=="__main__":
    df=make_seq_dataframe(doc,table_catalog_dict)
#     print(seq_dataframe.iloc[0:2,:])
    df=copy_seq_with_num_to_ID(df,"SEQUENCE","ID")
    
    # 清洗ID
    df=ID_clean(df, "ID")
    
    print(seq_dataframe)


# # 清洗

# In[14]:


if __name__=="__main__":
    df=make_seq_dataframe(doc,table_catalog_dict)
#     print(seq_dataframe.iloc[0:2,:])
    df=copy_seq_with_num_to_ID(df,"SEQUENCE","ID")
    # 清洗ID
    df=ID_clean(df,'ID')
    
    # 清洗name
    df=name_clean(df,'NAME')
    
    # 注意没有3字母序列, 所以不可清理3字母序列, 否则就没了

    # 清洗单字母序列
    df=seq_1letters_clean(df,'SEQUENCE')

    # 根据chaintype关键字清洗chaintype
    df=chaintype_clean(df,"CHAINTYPE")
    
    # 保存数据
    save_seq_to_csv(df,output_path,fname)


# # 封装

# In[15]:


def reconstruct_DOC(input_path,fname,output_path,table_catalog_dict):
    # 读取文件
    filename=os.path.join(input_path,fname)
    doc = Document(filename)
    
    df=make_seq_dataframe(doc,table_catalog_dict)
#     print(seq_dataframe.iloc[0:2,:])
    df=copy_seq_with_num_to_ID(df,"SEQUENCE","ID")
    # 清洗ID
    df=ID_clean(df,'ID')
    
    # 清洗name
    df=name_clean(df,'NAME')
    
    # 注意没有3字母序列, 所以不可清理3字母序列, 否则就没了
    
    # 清洗单字母序列
    df=seq_1letters_clean(df,'SEQUENCE')

    # 根据chaintype关键字清洗chaintype
    df=chaintype_clean(df,"CHAINTYPE")
    
    # 保存数据
    save_seq_to_csv(df,output_path,fname)
    
    return df
    

